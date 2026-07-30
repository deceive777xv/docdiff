"""Tests for app/services/ingest_service.py"""
from __future__ import annotations

import json
from unittest.mock import patch

import pytest


@pytest.fixture
def docx_file(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_heading("测试文档", level=1)
    doc.add_paragraph("这是测试内容。")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def db_conn(tmp_path):
    from app.db.schema import init_db
    conn = init_db(str(tmp_path))
    yield conn
    conn.close()


def test_ingest_returns_ids(tmp_path, docx_file, db_conn):
    """ingest_document returns (doc_id, version_id) as non-empty strings."""
    from app.services.ingest_service import ingest_document

    doc_id, version_id = ingest_document(
        db_conn, str(tmp_path), str(docx_file), embedder=None
    )

    assert isinstance(doc_id, str) and doc_id
    assert isinstance(version_id, str) and version_id


def test_ingest_duplicate_raises(tmp_path, docx_file, db_conn):
    """Ingesting the same file twice raises FileExistsError."""
    from app.services.ingest_service import ingest_document

    ingest_document(db_conn, str(tmp_path), str(docx_file), embedder=None)

    with pytest.raises(FileExistsError):
        ingest_document(db_conn, str(tmp_path), str(docx_file), embedder=None)


def test_ingest_creates_parsed_json(tmp_path, docx_file, db_conn):
    """After ingest, a parsed IR JSON file is created under data_dir/parsed/."""
    from app.services.ingest_service import ingest_document
    from pathlib import Path

    ingest_document(db_conn, str(tmp_path), str(docx_file), embedder=None)

    parsed_files = list((Path(tmp_path) / "parsed").glob("*.json"))
    assert len(parsed_files) >= 1


def test_ingest_inserts_chunks(tmp_path, docx_file, db_conn):
    """After ingest, get_chunks_by_version returns a non-empty list."""
    from app.services.ingest_service import ingest_document
    from app.db.chunk_repo import get_chunks_by_version

    _, version_id = ingest_document(
        db_conn, str(tmp_path), str(docx_file), embedder=None
    )

    chunks = get_chunks_by_version(db_conn, version_id)
    assert len(chunks) > 0


def test_ingest_persists_raw_and_indexes_normalized_ir(tmp_path, docx_file, db_conn):
    from app.core.normalization import NormalizationDepth
    from app.core.types import DocumentIR, Paragraph, ParseQualityReport, Section, Sentence
    from app.db import chunk_repo, document_repo
    from app.services.ingest_service import ingest_document

    raw = DocumentIR(
        doc_id="repair-doc",
        title="Repair",
        file_hash="repair-hash",
        sections=[
            Section(
                "s1",
                "**1.1 测试**",
                2,
                [Paragraph("p1", "正文。", [Sentence("正文。")], 1)],
            )
        ],
        plain_text="正文。",
    )
    quality = ParseQualityReport(quality_score=1.0, needs_ocr=False)

    with patch(
        "app.services.ingest_service.parse_document",
        return_value=(raw, quality),
    ):
        _, version_id = ingest_document(
            db_conn,
            str(tmp_path),
            str(docx_file),
            embedder=None,
            normalization_depth=NormalizationDepth.STANDARD,
        )

    version = document_repo.get_version_by_id(db_conn, version_id)
    normalized_path = tmp_path / "parsed" / "repair-doc.json"
    assert version["parsed_json_path"] == str(normalized_path)
    assert (tmp_path / "parsed" / "raw" / "repair-doc.json").exists()
    assert (tmp_path / "parsed" / "traces" / "repair-doc.structure.json").exists()
    normalized = json.loads(normalized_path.read_text(encoding="utf-8"))
    assert normalized["sections"][0]["title"] == "1.1 测试"
    chunks = chunk_repo.get_chunks_by_version(db_conn, version_id)
    assert chunks[0]["section_path"] == "1.1 测试"
