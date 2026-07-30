"""Tests for app/services/compare_service.py"""
from __future__ import annotations

from copy import deepcopy
import json
from pathlib import Path
from unittest.mock import MagicMock

import pytest


def _document_from_fixture(data):
    from app.core.types import DocumentIR, Paragraph, Section, Sentence

    return DocumentIR(
        doc_id=data["doc_id"],
        title=data["title"],
        file_hash=data["file_hash"],
        sections=[
            Section(
                section_id=section["section_id"],
                title=section["title"],
                level=section["level"],
                paragraphs=[
                    Paragraph(
                        paragraph_id=paragraph["paragraph_id"],
                        text=paragraph["text"],
                        sentences=[
                            Sentence(sentence["text"])
                            for sentence in paragraph["sentences"]
                        ],
                    )
                    for paragraph in section["paragraphs"]
                ],
            )
            for section in data["sections"]
        ],
        plain_text=data["plain_text"],
    )


def _sanitized_fixture():
    fixture_path = Path(__file__).parents[1] / "fixtures" / "cross_page_table_pair.json"
    payload = json.loads(fixture_path.read_text(encoding="utf-8"))
    return (
        fixture_path,
        _document_from_fixture(payload["baseline"]),
        _document_from_fixture(payload["target"]),
    )


def _normalized_pair_texts(pairs):
    return [
        (
            [paragraph.text for paragraph in pair.baseline_section.paragraphs]
            if pair.baseline_section else [],
            [paragraph.text for paragraph in pair.target_section.paragraphs]
            if pair.target_section else [],
        )
        for pair in pairs
    ]


@pytest.fixture
def two_docx_versions(tmp_path):
    from docx import Document

    doc1 = Document()
    doc1.add_heading("第一章 总则", level=1)
    doc1.add_paragraph("付款周期为30天。")
    p1 = tmp_path / "v1.docx"
    doc1.save(str(p1))

    doc2 = Document()
    doc2.add_heading("第一章 总则", level=1)
    doc2.add_paragraph("付款周期调整为60天。")
    p2 = tmp_path / "v2.docx"
    doc2.save(str(p2))

    return p1, p2


@pytest.fixture
def db_conn(tmp_path):
    from app.db.schema import init_db

    conn = init_db(str(tmp_path))
    yield conn
    conn.close()


def _make_mock_embedder():
    mock_embedder = MagicMock()
    mock_embedder.embed.side_effect = lambda texts: [
        [float(hash(t) % 1000) / 1000.0] * 8 for t in texts
    ]
    return mock_embedder


def _make_mock_provider():
    mock_provider = MagicMock()
    mock_provider.chat.return_value = (
        '{"diff_type": "实质修改", "risk_level": "high", "explanation": "金额变化"}'
    )
    return mock_provider


def test_run_compare_returns_diff_result(tmp_path, two_docx_versions, db_conn):
    """run_compare returns a DiffResult with at least one item."""
    from app.services.ingest_service import ingest_document
    from app.services.compare_service import run_compare
    from app.core.types import DiffResult

    p1, p2 = two_docx_versions
    mock_embedder = _make_mock_embedder()
    mock_provider = _make_mock_provider()

    _, v1_id = ingest_document(db_conn, str(tmp_path), str(p1), embedder=None)
    _, v2_id = ingest_document(db_conn, str(tmp_path), str(p2), embedder=None)

    result = run_compare(
        conn=db_conn,
        data_dir=str(tmp_path),
        baseline_version_id=v1_id,
        target_version_id=v2_id,
        embedder=mock_embedder,
        provider=mock_provider,
    )

    assert isinstance(result, DiffResult)
    assert len(result.items) >= 1


def test_compare_task_status_completed(tmp_path, two_docx_versions, db_conn):
    """After run_compare, the task record in DB has status 'completed'."""
    from app.services.ingest_service import ingest_document
    from app.services.compare_service import run_compare
    from app.db import compare_repo

    p1, p2 = two_docx_versions
    mock_embedder = _make_mock_embedder()
    mock_provider = _make_mock_provider()

    _, v1_id = ingest_document(db_conn, str(tmp_path), str(p1), embedder=None)
    _, v2_id = ingest_document(db_conn, str(tmp_path), str(p2), embedder=None)

    result = run_compare(
        conn=db_conn,
        data_dir=str(tmp_path),
        baseline_version_id=v1_id,
        target_version_id=v2_id,
        embedder=mock_embedder,
        provider=mock_provider,
    )

    task = compare_repo.get_task_by_id(db_conn, result.task_id)
    assert task is not None
    assert task["status"] == "completed"


def test_compare_detects_change(tmp_path, two_docx_versions, db_conn):
    """run_compare detects at least one diff item between the two versions."""
    from app.services.ingest_service import ingest_document
    from app.services.compare_service import run_compare

    p1, p2 = two_docx_versions
    mock_embedder = _make_mock_embedder()
    mock_provider = _make_mock_provider()

    _, v1_id = ingest_document(db_conn, str(tmp_path), str(p1), embedder=None)
    _, v2_id = ingest_document(db_conn, str(tmp_path), str(p2), embedder=None)

    result = run_compare(
        conn=db_conn,
        data_dir=str(tmp_path),
        baseline_version_id=v1_id,
        target_version_id=v2_id,
        embedder=mock_embedder,
        provider=mock_provider,
    )

    assert len(result.items) >= 1


def test_run_compare_marks_failed_and_reraises_when_result_publish_fails(tmp_path, monkeypatch):
    """Publishing must succeed before the synchronous service marks completion."""
    from app.core.types import DiffResult, DocumentIR
    from app.services import compare_service

    mock_ir = DocumentIR("doc", "Title", "hash")
    statuses: list[str] = []

    monkeypatch.setattr(
        compare_service.compare_repo,
        "create_compare_task",
        lambda *args, **kwargs: "task-sidecar",
    )
    monkeypatch.setattr(
        compare_service.compare_repo,
        "update_task_status",
        lambda conn, task_id, status, *args: statuses.append(status),
    )
    monkeypatch.setattr(compare_service.compare_repo, "insert_diff_items", lambda *args: None)
    monkeypatch.setattr(compare_service, "_load_ir", lambda *args: mock_ir)
    monkeypatch.setattr(compare_service, "align_sections", lambda *args: [])
    monkeypatch.setattr(compare_service, "match_paragraphs", lambda *args, **kwargs: [])
    monkeypatch.setattr(
        compare_service,
        "classify",
        lambda *args, **kwargs: DiffResult("task-sidecar", "ver-1", "ver-2", []),
    )
    monkeypatch.setattr(
        compare_service,
        "persist_compare_result",
        lambda *args: (_ for _ in ()).throw(OSError("result write failed")),
    )

    with pytest.raises(OSError, match="result write failed"):
        compare_service.run_compare(
            MagicMock(), str(tmp_path), "ver-1", "ver-2", MagicMock(), MagicMock()
        )

    assert statuses[-1] == "failed"
    assert "completed" not in statuses


def test_graph_and_service_consume_imported_irs_without_mutating_fixture(tmp_path, monkeypatch):
    """Both compare entry points pass imported IRs directly to matching."""
    from app.agent import compare_graph as graph_module
    from app.core.diff.result_storage import persist_compare_result
    from app.core.types import DiffResult
    from app.services import compare_service

    fixture_path, baseline, target = _sanitized_fixture()
    fixture_before = fixture_path.read_bytes()
    graph_texts = []
    service_texts = []
    provider = MagicMock()

    def load_ir(version_id, conn):
        return deepcopy(baseline if version_id == "ver-1" else target)

    def capture_graph_match(pairs, *args, **kwargs):
        graph_texts.append(_normalized_pair_texts(pairs))
        assert kwargs["baseline_document_title"] == baseline.title
        assert kwargs["target_document_title"] == target.title
        return []

    def capture_service_match(pairs, *args, **kwargs):
        service_texts.append(_normalized_pair_texts(pairs))
        return []

    monkeypatch.setattr(
        graph_module.compare_repo,
        "create_compare_task",
        lambda *args, **kwargs: "task-eq",
    )
    monkeypatch.setattr(
        graph_module.compare_repo,
        "update_task_status",
        lambda *args, **kwargs: None,
    )
    monkeypatch.setattr(
        graph_module.compare_repo,
        "insert_diff_items",
        lambda *args, **kwargs: None,
    )
    monkeypatch.setattr(graph_module, "_load_ir", load_ir)
    monkeypatch.setattr(graph_module, "match_paragraphs", capture_graph_match)
    monkeypatch.setattr(
        graph_module,
        "classify",
        lambda *args, **kwargs: DiffResult("task-eq", "ver-1", "ver-2", []),
    )
    monkeypatch.setattr(graph_module, "persist_compare_result", persist_compare_result)

    graph_state = graph_module.compare_graph.invoke(
        {
            "data_dir": str(tmp_path / "graph"),
            "baseline_version_id": "ver-1",
            "target_version_id": "ver-2",
            "provider": provider,
            "embedder": MagicMock(),
            "conn": MagicMock(),
        }
    )

    monkeypatch.setattr(
        compare_service.compare_repo,
        "create_compare_task",
        lambda *args, **kwargs: "task-eq",
    )
    monkeypatch.setattr(
        compare_service.compare_repo,
        "update_task_status",
        lambda *args, **kwargs: None,
    )
    monkeypatch.setattr(
        compare_service.compare_repo,
        "insert_diff_items",
        lambda *args, **kwargs: None,
    )
    monkeypatch.setattr(compare_service, "_load_ir", load_ir)
    monkeypatch.setattr(compare_service, "match_paragraphs", capture_service_match)
    monkeypatch.setattr(
        compare_service,
        "classify",
        lambda *args, **kwargs: DiffResult("task-eq", "ver-1", "ver-2", []),
    )
    monkeypatch.setattr(compare_service, "persist_compare_result", persist_compare_result)

    compare_service.run_compare(
        MagicMock(),
        str(tmp_path / "service"),
        "ver-1",
        "ver-2",
        MagicMock(),
        provider,
    )

    assert graph_state["status"] == "completed"
    assert graph_texts == service_texts
    assert fixture_path.read_bytes() == fixture_before
    assert not list(tmp_path.rglob("*.reconstruction.json"))


def test_service_does_not_call_table_reconstruction_during_compare(tmp_path, monkeypatch):
    """Table reconstruction belongs exclusively to import normalization."""
    from app.core.types import DiffResult
    from app.services import compare_service

    _, baseline, target = _sanitized_fixture()
    statuses: list[str] = []
    provider = MagicMock()

    def load_ir(version_id, conn):
        return deepcopy(baseline if version_id == "ver-1" else target)

    monkeypatch.setattr(
        compare_service.compare_repo,
        "create_compare_task",
        lambda *args, **kwargs: "task-provider-failure",
    )
    monkeypatch.setattr(
        compare_service.compare_repo,
        "update_task_status",
        lambda conn, task_id, status, *args: statuses.append(status),
    )
    monkeypatch.setattr(compare_service.compare_repo, "insert_diff_items", lambda *args: None)
    monkeypatch.setattr(compare_service, "_load_ir", load_ir)
    monkeypatch.setattr(compare_service, "match_paragraphs", lambda *args, **kwargs: [])
    monkeypatch.setattr(
        compare_service,
        "classify",
        lambda *args, **kwargs: DiffResult(
            "task-provider-failure", "ver-1", "ver-2", []
        ),
    )

    compare_service.run_compare(
        MagicMock(),
        str(tmp_path),
        "ver-1",
        "ver-2",
        MagicMock(),
        provider,
    )

    assert provider.chat.call_count == 0
    assert not list(tmp_path.rglob("*.reconstruction.json"))
    assert statuses[-1] == "completed"
