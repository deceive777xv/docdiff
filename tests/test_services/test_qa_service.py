"""Tests for app/services/qa_service.py"""
from __future__ import annotations

from unittest.mock import MagicMock

import pytest


@pytest.fixture
def docx_file(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("测试文档", level=1)
    doc.add_paragraph("这是测试内容。合同有效期为两年。")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def db_conn(tmp_path):
    from app.db.schema import init_db

    conn = init_db(str(tmp_path))
    yield conn
    conn.close()


def _make_mock_embedder():
    mock_embedder = MagicMock()
    mock_embedder.embed.side_effect = lambda texts: [[0.1] * 8 for _ in texts]
    return mock_embedder


def test_answer_no_versions_returns_message(tmp_path, db_conn):
    """answer() with empty current_version_ids returns a no-document message."""
    from app.services.qa_service import answer
    from app.core.types import RetrievalScope

    mock_embedder = _make_mock_embedder()
    mock_provider = MagicMock()

    answer_text, hits = answer(
        conn=db_conn,
        data_dir=str(tmp_path),
        question="这份合同的有效期是多久？",
        provider=mock_provider,
        embedder=mock_embedder,
        scope=RetrievalScope.CURRENT_DOC,
        current_version_ids=[],
        top_k=5,
    )

    assert isinstance(answer_text, str)
    assert len(answer_text) > 0
    assert hits == []


def test_answer_returns_text_and_hits(tmp_path, docx_file, db_conn):
    """answer() with an indexed doc returns a non-empty answer and chunk hits."""
    from app.services.ingest_service import ingest_document
    from app.services.qa_service import answer
    from app.core.types import RetrievalScope

    mock_embedder = _make_mock_embedder()
    mock_provider = MagicMock()
    mock_provider.chat.return_value = "合同有效期为两年。"

    # Ingest with embedder so FAISS index is built
    _, version_id = ingest_document(
        db_conn, str(tmp_path), str(docx_file), embedder=mock_embedder
    )

    answer_text, hits = answer(
        conn=db_conn,
        data_dir=str(tmp_path),
        question="合同有效期是多久？",
        provider=mock_provider,
        embedder=mock_embedder,
        scope=RetrievalScope.CURRENT_DOC,
        current_version_ids=[version_id],
        top_k=3,
    )

    assert isinstance(answer_text, str)
    assert isinstance(hits, list)
    assert len(hits) > 0


def test_answer_standard_lib_scope(tmp_path, docx_file, db_conn):
    """answer() with STANDARD_LIB scope retrieves from standard-source docs."""
    from app.services.ingest_service import ingest_document
    from app.services.qa_service import answer
    from app.core.types import RetrievalScope

    mock_embedder = _make_mock_embedder()
    mock_provider = MagicMock()
    mock_provider.chat.return_value = "标准库回答内容。"

    # Ingest as a standard document with FAISS index
    ingest_document(
        db_conn,
        str(tmp_path),
        str(docx_file),
        source_type="standard",
        embedder=mock_embedder,
    )

    answer_text, hits = answer(
        conn=db_conn,
        data_dir=str(tmp_path),
        question="有效期是多少年？",
        provider=mock_provider,
        embedder=mock_embedder,
        scope=RetrievalScope.STANDARD_LIB,
        top_k=3,
    )

    assert isinstance(answer_text, str)
    assert len(hits) > 0
