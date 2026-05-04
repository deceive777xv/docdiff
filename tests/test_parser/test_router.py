"""Tests for app/core/parser/router.py"""
from __future__ import annotations
import pytest
from pathlib import Path

from app.core.types import DocumentIR


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    from docx import Document
    doc = Document()
    doc.add_heading("第一章 总则", level=1)
    doc.add_paragraph("本协议自双方签署之日起生效。")
    doc.add_paragraph("甲方应在签署后30日内完成付款。")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_pdf(tmp_path) -> Path:
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "第一章 总则", fontsize=18)
    page.insert_text((72, 120), "本协议自双方签署之日起生效。", fontsize=12)
    path = tmp_path / "sample.pdf"
    doc.save(str(path))
    doc.close()
    return path


def test_routes_docx(sample_docx):
    """parse_document routes .docx to docx_extractor and returns DocumentIR."""
    from app.core.parser.router import parse_document
    ir, report = parse_document(str(sample_docx), mode="fast")
    assert isinstance(ir, DocumentIR)


def test_routes_pdf_fast(sample_pdf):
    """parse_document with mode='fast' routes .pdf to PyMuPDF and returns DocumentIR."""
    from app.core.parser.router import parse_document
    ir, report = parse_document(str(sample_pdf), mode="fast")
    assert isinstance(ir, DocumentIR)


def test_routes_pdf_standard_docling_unavailable(sample_pdf, monkeypatch):
    """When docling is unavailable, standard mode falls back to PyMuPDF."""
    import app.core.parser.router as router_module

    # Monkeypatch is_available inside the router module
    monkeypatch.setattr(
        "app.core.parser.docling_adapter.is_available",
        lambda: False,
    )

    from app.core.parser.router import parse_document
    ir, report = parse_document(str(sample_pdf), mode="standard")
    assert isinstance(ir, DocumentIR)


def test_unsupported_format_raises(tmp_path):
    """parse_document raises ValueError for unsupported file extensions."""
    from app.core.parser.router import parse_document
    txt_file = tmp_path / "notes.txt"
    txt_file.write_text("some content")
    with pytest.raises(ValueError, match="Unsupported file format"):
        parse_document(str(txt_file))
