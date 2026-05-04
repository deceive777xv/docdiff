"""Tests for app/core/parser/docx_extractor.py"""
import pytest
from docx import Document
from pathlib import Path

from app.core.types import DocumentIR
from app.core.parser.docx_extractor import extract


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    doc = Document()
    doc.add_heading("第一章 总则", level=1)
    doc.add_paragraph("本协议自双方签署之日起生效。")
    doc.add_paragraph("甲方应在签署后30日内完成付款。")
    doc.add_heading("第二章 违约责任", level=2)
    doc.add_paragraph("如有违约，违约方需支付合同金额的20%作为违约金。")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


def test_extract_returns_document_ir(sample_docx):
    ir, report = extract(str(sample_docx))
    assert isinstance(ir, DocumentIR)


def test_title_from_first_heading(sample_docx):
    ir, _ = extract(str(sample_docx))
    assert ir.title == "第一章 总则"


def test_sections_extracted(sample_docx):
    ir, _ = extract(str(sample_docx))
    assert len(ir.sections) >= 2
    for section in ir.sections:
        assert section.title  # non-empty title


def test_paragraphs_in_sections(sample_docx):
    ir, _ = extract(str(sample_docx))
    all_paragraphs = [p for sec in ir.sections for p in sec.paragraphs]
    assert len(all_paragraphs) > 0
    for p in all_paragraphs:
        assert p.text  # non-empty text


def test_quality_report_no_ocr(sample_docx):
    _, report = extract(str(sample_docx))
    assert report.needs_ocr is False
    assert report.quality_score > 0.5
