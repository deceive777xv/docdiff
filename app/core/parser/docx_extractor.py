"""Extract DocumentIR from .docx files using python-docx."""
from __future__ import annotations
import re
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.core.types import DocumentIR, Paragraph, ParseQualityReport, Section, Sentence
from app.core.utils import file_hash


def _heading_level(para) -> int | None:
    """Return 1/2/3 if paragraph is a heading, else None."""
    style_name = para.style.name.lower()
    for level in (1, 2, 3):
        if f"heading {level}" in style_name:
            return level
    return None


def _split_sentences(text: str) -> list[Sentence]:
    parts = re.split(r'(?<=[。！？.!?])\s*', text.strip())
    return [Sentence(text=p.strip()) for p in parts if p.strip()]


def extract(file_path: str) -> tuple[DocumentIR, ParseQualityReport]:
    """
    Parse a .docx file into DocumentIR.
    Returns (DocumentIR, ParseQualityReport).
    """
    path = Path(file_path)
    doc = Document(str(path))
    doc_hash = file_hash(path)
    doc_id = str(uuid.uuid4())

    title = path.stem
    sections: list[Section] = []
    current_section: Section | None = None
    para_count = 0
    empty_count = 0

    # Default section for content before any heading
    default_section = Section(
        section_id=str(uuid.uuid4()),
        title="正文",
        level=1,
    )

    for para in doc.paragraphs:
        text = para.text.strip()
        level = _heading_level(para)

        if level is not None and text:
            # Start a new section at this heading
            if level == 1 and text and title == path.stem:
                title = text  # Use first H1 as document title
            current_section = Section(
                section_id=str(uuid.uuid4()),
                title=text,
                level=level,
            )
            sections.append(current_section)
        elif text:
            para_count += 1
            target = current_section if current_section is not None else default_section
            if target is default_section and default_section not in sections:
                sections.insert(0, default_section)
            p = Paragraph(
                paragraph_id=str(uuid.uuid4()),
                text=text,
                sentences=_split_sentences(text),
            )
            target.paragraphs.append(p)
        else:
            empty_count += 1

    plain_text = "\n".join(
        p.text
        for sec in sections
        for p in sec.paragraphs
    )

    total = para_count + empty_count
    quality_score = 1.0 if total == 0 else max(0.0, 1.0 - empty_count / max(total, 1))

    ir = DocumentIR(
        doc_id=doc_id,
        title=title,
        file_hash=doc_hash,
        sections=sections,
        plain_text=plain_text,
    )
    report = ParseQualityReport(
        quality_score=round(quality_score, 2),
        needs_ocr=False,
    )
    return ir, report
