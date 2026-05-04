"""Export diff results to Word (.docx) or HTML."""
from __future__ import annotations

import html as html_mod
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from app.core.types import DiffResult

_RISK_LABELS = {"high": "高风险", "medium": "中风险", "low": "低风险"}

_DIFF_COLORS_HEX: dict[str, str] = {
    "新增":     "22c55e",
    "删减":     "ef4444",
    "微调":     "eab308",
    "实质修改": "f97316",
    "重写":     "a855f7",
    "格式变化": "9ca3af",
}


def _hex_to_rgb(h: str) -> tuple[int, int, int]:
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def export_docx(result: DiffResult, output_path: str) -> None:
    """Export DiffResult to a Word document at output_path."""
    doc = Document()

    doc.add_heading("文档对比报告", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"差异总数：{len(result.items)}")

    counts = Counter(item.diff_type for item in result.items)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading"
    hdr = table.rows[0].cells
    hdr[0].text = "差异类型"
    hdr[1].text = "数量"
    for dtype, count in counts.most_common():
        row = table.add_row().cells
        row[0].text = dtype
        row[1].text = str(count)

    doc.add_page_break()
    doc.add_heading("差异详情", 1)

    for i, item in enumerate(result.items, 1):
        doc.add_heading(f"{i}. [{item.diff_type}] {item.section_path}", 2)
        risk_label = _RISK_LABELS.get(item.risk_level, item.risk_level)
        doc.add_paragraph(f"风险等级：{risk_label}  |  相似度：{item.similarity_score:.3f}")
        if item.baseline_text:
            p = doc.add_paragraph()
            run = p.add_run("基准：")
            run.bold = True
            p.add_run(item.baseline_text[:400])
        if item.target_text:
            p = doc.add_paragraph()
            run = p.add_run("目标：")
            run.bold = True
            p.add_run(item.target_text[:400])
        if item.explanation:
            p = doc.add_paragraph(f"说明：{item.explanation}")
            p.runs[0].font.size = Pt(10)
            r, g, b = _hex_to_rgb("6b7280")
            p.runs[0].font.color.rgb = RGBColor(r, g, b)

    doc.save(output_path)


def export_html(result: DiffResult, output_path: str) -> None:
    """Export DiffResult to a standalone HTML report at output_path."""
    counts = Counter(item.diff_type for item in result.items)
    stats_rows = "".join(
        f"<tr><td>{html_mod.escape(dt)}</td><td>{cnt}</td></tr>"
        for dt, cnt in counts.most_common()
    )

    items_html_parts: list[str] = []
    for item in result.items:
        color = _DIFF_COLORS_HEX.get(item.diff_type, "9ca3af")
        risk = _RISK_LABELS.get(item.risk_level, item.risk_level)
        section = html_mod.escape(item.section_path)
        b_text = html_mod.escape(item.baseline_text[:400]) if item.baseline_text else ""
        t_text = html_mod.escape(item.target_text[:400]) if item.target_text else ""
        exp = html_mod.escape(item.explanation) if item.explanation else ""

        baseline_block = (
            f'<div class="baseline"><strong>基准：</strong>{b_text}</div>' if b_text else ""
        )
        target_block = (
            f'<div class="target"><strong>目标：</strong>{t_text}</div>' if t_text else ""
        )
        explain_block = (
            f'<div class="explain">说明：{exp}</div>' if exp else ""
        )

        items_html_parts.append(f"""
    <div class="diff-card" style="border-left:4px solid #{color}">
      <div class="header">
        <span class="badge" style="background:#{color}">{html_mod.escape(item.diff_type)}</span>
        <span class="risk">{html_mod.escape(risk)}</span>
        <span class="section">{section}</span>
        <span class="sim">相似度：{item.similarity_score:.3f}</span>
      </div>
      {baseline_block}
      {target_block}
      {explain_block}
    </div>""")

    items_html = "\n".join(items_html_parts)
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")

    html_content = f"""<!DOCTYPE html>
<html lang="zh">
<head>
  <meta charset="UTF-8">
  <title>文档对比报告</title>
  <style>
    body {{ font-family: "Segoe UI", "Microsoft YaHei", sans-serif; max-width: 1000px;
           margin: 40px auto; color: #1f2937; background: #f9fafb; }}
    h1 {{ color: #1e3a5f; margin-bottom: 4px; }}
    .meta {{ color: #6b7280; font-size: 13px; margin-bottom: 24px; }}
    h2 {{ color: #374151; border-bottom: 2px solid #e5e7eb; padding-bottom: 6px; }}
    table {{ border-collapse: collapse; margin: 16px 0; background: white;
             border-radius: 8px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,.1); }}
    td, th {{ border: 1px solid #e5e7eb; padding: 8px 16px; text-align: left; }}
    th {{ background: #f3f4f6; font-weight: 600; }}
    .diff-card {{ background: white; border: 1px solid #e5e7eb; border-radius: 8px;
                  padding: 12px 16px; margin: 12px 0;
                  box-shadow: 0 1px 2px rgba(0,0,0,.05); }}
    .header {{ display: flex; align-items: center; gap: 8px; margin-bottom: 8px; }}
    .badge {{ color: white; border-radius: 4px; padding: 2px 8px;
              font-size: 12px; font-weight: bold; }}
    .risk {{ font-size: 12px; color: #6b7280; }}
    .section {{ font-size: 13px; font-weight: 600; }}
    .sim {{ font-size: 11px; color: #9ca3af; margin-left: auto; }}
    .baseline {{ background: #fef2f2; border-radius: 4px; padding: 6px 10px;
                 margin: 6px 0; font-size: 13px; line-height: 1.6; }}
    .target {{ background: #f0fdf4; border-radius: 4px; padding: 6px 10px;
               margin: 6px 0; font-size: 13px; line-height: 1.6; }}
    .explain {{ color: #6b7280; font-size: 12px; margin-top: 6px; }}
  </style>
</head>
<body>
  <h1>文档对比报告</h1>
  <p class="meta">生成时间：{generated_at} &nbsp;|&nbsp; 差异总数：{len(result.items)}</p>
  <h2>差异统计</h2>
  <table>
    <tr><th>差异类型</th><th>数量</th></tr>
    {stats_rows}
  </table>
  <h2>差异详情</h2>
  {items_html}
</body>
</html>"""

    Path(output_path).write_text(html_content, encoding="utf-8")
